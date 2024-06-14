# encoding:utf-8
# pip install pydocx python-docx
import time
import itertools
from PIL import Image
from docx import Document
from pydocx.export.html import PyDocXHTMLExporter, HtmlTag
from itertools import chain


jss = '''
        document.addEventListener('DOMContentLoaded', function () {
            let dataForms = document.getElementsByTagName('form')
            for (let i = 0; i < dataForms.length; i++) {
                dataForms[i].addEventListener('submit', function (e) {
                    e.preventDefault();
                    addRow(i, dataForms[i].getElementsByTagName("input").length);
                });
            }
            document.querySelectorAll('tbody tr td').forEach(function (cell) {
                if (cell.innerHTML.length === 0) {
                    cell.addEventListener('dblclick', function () {
                        makeEditable(cell);
                    });
                }
            });
            document.addEventListener('input', function (e) {
                e.target.style.width = e.target.scrollWidth + 'px';
            });
            document.addEventListener('keydown', function (e) {
                if (e.key === 'Enter') {
                    if (e.target.tagName === "INPUT" && e.target.className === "pydocx-underline") {
                        nextInput(e.target)
                    }
                }
            });
            window.addEventListener('message', function (e) {
                if (e.data === "save") {
                    e.source.postMessage(inputData(), '*');
                } else {
                    let datas = e.data.inputs;
                    let tables = e.data.tables
                    let inputs = document.getElementsByClassName("pydocx-underline");
                    inputs[0].focus();
                    if (datas.length === inputs.length) {
                        for (let i = 0; i < inputs.length; i++) {
                            inputs[i].type = datas[i].type
                            inputs[i].value = datas[i].value
                        }
                    }
                    if (typeof(exp) != "undefined"){
                        let tbodys = document.getElementsByTagName("tbody");
                        if (tables.length === tbodys.length) {
                            for (let i = 0; i < tables.length; i++) {
                                let rcount = tbodys[i].childElementCount;
                                for (let j = 0; j < tables[i].length; j++) {
                                    if (j < rcount) {
                                        for (let k = 0; k < tables[i][j].length; k++) {
                                            tbodys[i].children[j].cells[k].innerHTML = tables[i][j][k]
                                        }
                                    } else {
                                        let row = tbodys[i].insertRow();
                                        for (let k = 0; k < tables[i][j].length; k++) {
                                            let cell = row.insertCell(k);
                                            cell.textContent = tables[i][j][k]
                                            makeEditable(cell);
                                        }
                                    }
                                }
                            }
                        }
                    }
                }
            });
        });
        function nextInput(thisInput) {
            let inputs = document.getElementsByClassName("pydocx-underline");
            for (let i = 0; i < inputs.length+1; i++) {
                if (thisInput == inputs[i]) {
                    if (i == inputs.length - 1) {
                        inputs[i].focus();
                    } else {
                        inputs[i + 1].focus();
                    }
                    break;
                }
                if (i === inputs.length){ 
                    inputs[0].focus();
                }
            }
        };
        function inputData() {
            let inputs = [];
            Array.from(document.getElementsByClassName("pydocx-underline")).forEach(element => {
                inputs.push(element.value)
            });
            let tables = [];
            let tbodys = document.getElementsByTagName("tbody");
            for (let k = 0; k < tbodys.length; k++) {
                let tbody = [];
                for (let i = 0; i < tbodys[k].childElementCount; i++) {
                    let row = tbodys[0].children[i];
                    let rows = [];
                    for (let j = 0; j < row.childElementCount; j++) {
                        rows.push(row.cells[j].innerHTML);
                    }
                    tbody.push(rows);
                }
                tables.push(tbody);
            }
            return { inputs, tables }
        };
        function addRow(i, n) {
            let tableBody = document.getElementsByTagName('tbody')[i];
            let row = tableBody.insertRow();
            for (let i = 0; i < n; i++) {
                let cell = row.insertCell(i);
                cell.textContent = document.getElementById('form' + (i + 1)).value;
                makeEditable(cell);
            }
            let dataForms = document.getElementsByTagName('form')
            for (let i = 0; i < dataForms.length; i++) {
                dataForms[i].reset();
            }
        };
        function makeEditable(cell) {
            cell.contentEditable = true;
            cell.focus();
        };
'''


class item():
    data = None
    required = True
    date = None
    table = None

    def __init__(self, detail=""):
        self.detail = detail


class docx2html(PyDocXHTMLExporter):
    colums = []
    inputDatas = []
    inputData = []

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)

    def head(self):
        tag = HtmlTag('head')
        results = chain(self.meta(), self.style(), self.js())
        return tag.apply(results)

    def js(self):
        tag = HtmlTag('script')
        return tag.apply(jss)

    # 表格
    def get_table_tag(self, table):
        return HtmlTag('table', **{
            'width': '100%',
            "border": '1',
        })

    def export_table(self, table):
        results = list(itertools.chain.from_iterable([super().export_table(table)]))
        for i, result in enumerate(results):
            if i == 1:
                yield '<thead>'
            if len(results) == i+2:
                yield '</tbody>'
            else:
                yield result

        yield '<form class="dataForm">'
        for i, _ in enumerate(self.colums[0]):
            yield '<input id="form%d" class="table-input"></input>' % (i+1,)
        yield '<button type="submit">添加</button>'
        yield '</form>'

    def export_table_row(self, table_row):
        colums = []
        results = super(PyDocXHTMLExporter, self).export_table_row(table_row)
        res = []
        for result in results:
            if isinstance(result, str):
                colums.append(result)
            res.append(result)
        if colums:
            i = 0
            for v in res:
                if isinstance(v, HtmlTag):
                    if v.tag == "td":
                        if not v.closed:
                            i += 1
                            v.attrs["field"] = 'form%d' % i
                        yield v
                else:
                    yield v
            if colums not in self.colums:
                self.colums.append(colums)
            if len(colums) == len(table_row.cells):
                yield '</thead><tbody>'
            else:
                yield '</tr>'
            yield '<tr>'

    # 输入
    def export_run_property_underline(self, run, results):
        self.reset_data()
        attrs = {
            'class': 'pydocx-underline',
        }
        tag = HtmlTag('input', **attrs)
        return tag.apply(results)

    def reset_data(self):
        if len(self.inputData) > 1:
            data = self.inputData[-1]
            if len(data) < 2 and len(self.inputData) > 2:
                data = self.inputData[-2] + data
            self.inputDatas.append(data)
            self.inputData = []
        elif len(self.inputDatas) > 0:
            self.inputDatas.append("")

    # 图片缩放
    def get_image_tag(self, image, width=None, height=None, rotate=None):
        attrs = {
            'width': width,
            'height': height
        }
        c = rotate and rotate in (270, 90)
        ic = Image.open(image.stream)
        if (ic.width <= 595 and not c) or (ic.height <= 842 and c):
            image_src = self.get_image_source(image)
            if rotate:
                attrs['style'] = 'transform: rotate(%sdeg);' % rotate
        else:
            import base64
            from io import BytesIO
            if c:
                ic = ic.rotate(360-rotate, expand=True)
                attrs["width"], attrs["height"] = height, width
            ic = ic.resize((842, int(ic.height * 842 / ic.width)))
            buffer = BytesIO()
            ic.save(buffer, format="JPEG")
            image_src = 'data:image/JPEG;base64,{data}'.format(data=base64.b64encode(buffer.getvalue()).decode('utf-8'))
        if image_src:
            return HtmlTag(
                'img',
                allow_self_closing=True,
                allow_whitespace=True,
                src=image_src,
                **attrs
            )

    # 导出html
    def export(self):
        re = []
        for result in super(PyDocXHTMLExporter, self).export():
            if isinstance(result, HtmlTag):
                re.append(result.to_html())
            else:
                self.inputData.append(result)
                re.append(result)
        return ''.join(re)

    # 模版数据替换
    def save(self, paths="", datas=[], tables=[]):
        paths = paths or self.path.replace(".docx", "_%d.docx" % time.time())
        doc, i = Document(self.path), 0
        for para in doc.paragraphs:
            for run in para.runs:
                if i >= len(datas):
                    continue
                if run.underline:
                    run.text = datas[i]
                    i += 1

        for i, table in enumerate(tables):
            rn = len(doc.tables[i].rows)-1  # 第一行标题
            while rn < len(table):
                doc.tables[i].add_row()
                rn += 1
            for j, row in enumerate(doc.tables[i].rows):
                if j > 0:
                    for k, cell in enumerate(row.cells):
                        cell.text = table[j-1][k]
        doc.save(paths)
        return paths


if __name__ == "__main__":
    docs = docx2html('用户协议.docx')
    html = docs.export()
    with open("res.html", "w", encoding="utf-8") as f:
        f.write(html)
    # docs.save("", ["替换数据" for i in range(30)], [
    #     [["1", "2", "3\n5", "4", "5"] for i in range(6)]
    # ]
    # )
    print(docs.colums)
    print(docs.inputDatas)
